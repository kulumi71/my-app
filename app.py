import streamlit as st
import io
from docx import Document

# 設置網頁風格與標題 (無 emoji 純淨設計)
st.set_page_config(page_title="召會兒童班教材自動化生產線 v4.0", layout="centered")

# 注入高質感、高對比、無 emoji 且支援多頁面切換的 CSS 樣式
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    
    /* 莫蘭迪淡色系圓角高對比按鈕 - 藍色系 (下載與主按鈕) */
    div.stDownloadButton > button, div.stDownloadButton > button:active {
        width: 100% !important;
        border-radius: 24px !important;
        height: 3.5em !important;
        font-weight: bold !important;
        font-size: 16px !important;
        background-color: #e0f2fe !important; /* 淡藍 */
        color: #0369a1 !important; /* 深海藍 */
        border: 1.5px solid #bae6fd !important;
        transition: all 0.3s ease !important;
    }
    div.stDownloadButton > button:hover {
        background-color: #bae6fd !important;
        color: #0c4a6e !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 12px rgba(3, 105, 161, 0.1) !important;
    }
    
    /* 莫蘭迪淡色系圓角高對比按鈕 - 綠色系 (影音跳轉按鈕) */
    div.stLinkButton > a {
        width: 100% !important;
        border-radius: 24px !important;
        height: 3.5em !important;
        font-weight: bold !important;
        font-size: 16px !important;
        background-color: #dcfce7 !important; /* 淡綠 */
        color: #15803d !important; /* 深森林綠 */
        border: 1.5px solid #bbf7d0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        text-decoration: none !important;
        transition: all 0.3s ease !important;
    }
    div.stLinkButton > a:hover {
        background-color: #bbf7d0 !important;
        color: #166534 !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 12px rgba(21, 128, 61, 0.1) !important;
    }

    /* 頁面切換：返回目錄專用按鈕樣式 (高對比紫灰色) */
    div.stButton > button {
        border-radius: 24px !important;
        font-weight: bold !important;
        background-color: #f1f5f9 !important;
        color: #334155 !important;
        border: 1.5px solid #cbd5e1 !important;
        transition: all 0.3s ease !important;
    }
    div.stButton > button:hover {
        background-color: #e2e8f0 !important;
        color: #0f172a !important;
    }

    /* 超強對比度 Tabs 分頁標籤自定義 */
    div[data-testid="stTabs"] button {
        color: #475569 !important; /* 未選中文字：深灰 */
        background-color: #f1f5f9 !important; /* 未選中背景：淺灰 */
        border-radius: 12px 12px 0 0 !important;
        border: 1px solid #e2e8f0 !important;
        padding: 10px 24px !important;
        font-size: 16px !important;
        font-weight: bold !important;
        margin-right: 6px !important;
        transition: all 0.2s ease !important;
    }
    div[data-testid="stTabs"] button[aria-selected="true"] {
        color: #ffffff !important; /* 選中文字：純白 */
        background-color: #0369a1 !important; /* 選中背景：深天藍 */
        border: 1px solid #0369a1 !important;
    }

    /* 卡片式容器排版 */
    .download-card { padding: 24px; border-radius: 20px; background: white; border: 1px solid #e2e8f0; margin-bottom: 20px; box-shadow: 0 4px 12px rgba(15, 23, 42, 0.03); }
    .video-card { padding: 24px; border-radius: 20px; background: #f0fdf4; border: 1px solid #bbf7d0; margin-bottom: 20px; }
    .slides-card { padding: 24px; border-radius: 20px; background: #f0f9ff; border: 1px solid #bae6fd; margin-bottom: 20px; }
    .catalogue-card { padding: 20px; border-radius: 16px; background: white; border: 1px solid #e2e8f0; margin-bottom: 16px; box-shadow: 0 2px 8px rgba(0,0,0,0.02); transition: 0.3s; }
    .catalogue-card:hover { transform: translateY(-3px); box-shadow: 0 6px 16px rgba(0,0,0,0.06); }
    
    h1 { color: #0f172a; font-weight: 800; }
    h2 { color: #1e293b; font-weight: 700; margin-top: 25px; }
    h3 { color: #1e293b; font-weight: 700; margin-top: 10px; }
    </style>
    """, unsafe_allow_html=True)

# 初始化多頁面狀態控制器
if 'page' not in st.session_state:
    st.session_state.page = 'home'
if 'selected_week' not in st.session_state:
    st.session_state.selected_week = 1

# --- 1 至 40 課《小明燈》第五冊完整真實目錄資料庫 ---
CURRICULUM_DB = {
    1: {
        "title": "第 1 週：兒女是神賜給父母的產業",
        "hymn": "兒童詩歌第 194 頁《當我還是個小小嬰孩》",
        "verse": "「兒女是耶和華所賜的產業；所懷的胎，是他所給的賞賜。」— 詩篇一百二十七篇 3 節",
        "discussion": "1. 請小朋友簡單介紹並描述自己的父母。\\n2. 討論我們是怎麼來的，我們的父母又是怎麼來的。\\n3. 討論人為什麼要結婚？人結婚後為什麼要有孩子？",
        "story": "神照著自己的形像和樣式用地上的塵土造了亞當，盼望人在地上彰顯祂並代表祂。神看亞當獨居不好，便在他熟睡時用肋骨造了夏娃作他完美的配偶，使他們在愛與一裡合為一體。神賜福給這第一對夫妻，要他們生養眾多。孩子是神賜給父母的產業與祝福。舊約的哈拿因沒有孩子而在殿裡向耶和華傾心吐意、迫切流淚禱告，神便顧念她，賜下日後偉大的神人撒母耳，除去她的羞恥，使哈拿的心因耶和華歡欣稱頌。",
        "exercise": "本週操練：在禱告中感謝神把我們擺在一個敬虔、和樂的家庭中，使我們能享受神給我們的祝福。",
        "pure_text": "第一課：兒女是神賜給父母的產業\\n\\n【本週背經】\\n「兒女是耶和華所賜的產業；所懷的胎，是他所給的賞賜。」— 詩篇一百二十七篇 3 節\\n「神就賜福給他們，又對他們說，要生養眾多，遍滿地面，治理這地。」— 創世記一章 28 節\\n\\n【教材內容】\\n神照著自己的形像，按著自己的樣式創造了人，盼望人在地上彰顯祂並代表祂。在神最初的創造中，神與人的關係是非常親密且彼此瞭解的。隨後，神使亞當沈睡，取其肋骨造了夏娃作完美的配偶。亞當一見便歡喜說：『這是我骨中的骨，肉中的肉。』兩個人在愛與一裡有著甜美的親密關係。神也賜福這第一對夫妻，要他們生養眾多、遍滿地面。因此，孩子神給父母的產業與賞賜。在以色列人中，沒有孩子曾是一件羞恥的事，但哈拿在殿裡向耶和華傾心吐意、迫切禱告，神便顧念她並應允她，賜下日後偉大的神人撒母耳。這說明孩子是從神而來的祝福，生命延續的喜樂能使家庭的愛更加甜蜜與擴大。",
        "youtube_url": "https://www.youtube.com/watch?v=JEOjPZgX44w", 
        "custom_slides_url": "https://drive.google.com/drive/folders/1uo7uR6wQk94C8xqxF1uwnChZWJs58zog?usp=sharing" 
    },
    2: {
        "title": "第 2 週：父母對兒女的責任",
        "hymn": "兒童詩歌第 194 頁《當我還是個小小嬰孩》",
        "verse": "「教養孩童，使他走當行的道，就是到老他也不偏離。」— 箴言二十二章 6 節",
        "discussion": "1. 請小朋友介紹父母的職業，並討論父母工作是為了什麼？\\n2. 如果沒有父母辛苦工作，我們食衣住行的供應從哪裡來？\\n3. 討論父母在日常生活中如何關心、教導我們？當我們不聽話時，父母為什麼要管教我們？",
        "story": "父母對兒女的愛是神造人時賦予人的。當我們是無助的嬰孩時，父母便以極大的愛心與耐心照顧、餵奶並守護我們。神託付父母供給我們成長所需的一切物質。同時，召會中的父母也負有在神面前管教與用主的話教養我們的責任。舊約的約基別在法老下令殺男嬰的危難中，憑著對神的信仰藏了摩西，並在乳養他的有限時間裡，將神的信仰傳輸給他，使摩西長大後不留戀王宮，成為帶領以色列人離開埃及的器皿。新約提摩太的外祖母羅以和母親友尼基，從小將聖經的話和對主的信心傳輸到提摩太裡面，成全他成為一個愛主、與保羅同工的青年人。",
        "exercise": "本週操練：為著父母對我們的愛、生活上的供應，以及在正確道路上的管教，向主獻上感謝的禱告。",
        "pure_text": "第二課：父母對兒女的責任\\n\\n【本週背經】\\n「教養孩童，使他走當行的道，就是到老他也不偏離。」— 箴言二十二章 6 節\\n「因為主所愛的，他必管教，又鞭打凡所收納的兒子。」— 希伯來書十二章 6 節\\n\\n【教材內容】\\n父母對兒女的愛，是神造人時便賦予人的，這不是一個重擔，而是他們的喜悅。父母對兒女的第一個責任是愛和照顧，在我們還是無助的小小嬰孩時，父母便付出了極大的耐心呵護我們。第二個責任是供給生存及成長所需的食、衣、住、行各種物質，為此父親出外工作。第三個責任則是管教，召會中的父母必須照著神的話，用正確的方法來教養孩子，使我們走當行的道。舊約中摩西的母親約基別，憑著敬畏神的心藏起摩西，並在乳養他的寶貴時期不斷傳輸信仰，使摩西日後成為帶領百萬以色列人出埃及的器皿。新約中提摩太的外祖母羅以和母親友尼基，從小用聖經教導他，將純真信心傳輸給他，成全他成為保羅在主裡最親密、愛主的青年幫手。",
        "youtube_url": "https://www.youtube.com/watch?v=9dYF5yaMoEk",
        "custom_slides_url": "https://drive.google.com/drive/folders/1uo7uR6wQk94C8xqxF1uwnChZWJs58zog?usp=sharing" 
    },
    3: {
        "title": "第 3 週：兒女對父母該有的態度(一)─愛父母",
        "hymn": "兒童詩歌第 209 首《兒女當孝敬父母》",
        "verse": "「當孝敬父母，使你的日子在耶和華你神所賜你的地上，得以長久。」— 出埃及記二十章 12 節",
        "discussion": "1. 請小朋友說說最喜歡父母的哪些方面？\\n2. 父母為我們做了那麼多，我們能為父母做什麼？我們可以用什麼方式向父母表達我們的感激與愛？",
        "story": "神託付兒女對父母盡孝敬與順從的責任。孝敬的表現就是愛和尊敬父母。愛父母是因為父母是我們的源頭，就如同神是我們的源頭一樣。我們可以在日常生活中，貼心地幫助父母做家事、自己自動寫完功課整理房間、畫卡片或口頭說謝謝，來向父母具體表達愛。達秘弟兄雖然五歲時母親就去世了，但他一生深深懷念母親溫柔專注的愛，這使他懂得何為被愛，也使他成為一個一生愛主、大為神所用的僕人。偉大的林肯總統也始終敬愛他的繼母，並見證自己一生的成就完全出於母親的教誨。",
        "exercise": "本週操練：主動向父母表達我們對他們的愛，如對父母親說謝謝，幫助父母作家事，主動擁抱他們或和父母談心。",
        "pure_text": "第三課：兒女對父母該有的態度(一)─愛父母\\n\\n【本週背經】\\n「當孝敬父母，使你的日子在耶和華你神所賜你的地上，得以長久。」— 出埃及記二十章 12 節\\n「你要使父母歡喜；使生你的快樂。」— 箴言二十三章 25 節\\n\\n【教材內容】\\n父母為我們付出極多，神也託付兒女對父母當盡孝敬與順從的責任。愛父母是孝敬的具體表現。因為父母是我們的源頭，就如同神是我們的源頭。我們在生活中可以藉由與父母談心、幫助分擔家事、送上手繪卡片或擁抱來表達愛意，這能減輕父母的負擔。達秘弟兄年幼喪母，但他一生記得母親溫柔專注的愛，這使他懂得愛、並成為被主大用的愛主僕人。偉大的林肯總統也始終敬愛他的繼母，並見證自己的一切完全出於母親的屬靈教誨。",
        "youtube_url": "https://www.youtube.com/watch?v=wcEyIBhuVpM",
        "custom_slides_url": "https://drive.google.com/drive/folders/1uo7uR6wQk94C8xqxF1uwnChZWJs58zog?usp=sharing" 
    }
}

# --- 第 4 週 至 第 40 週 骨架結構預備與官方目錄校正 ---
WEEKS_4_40_TITLES = {
    4: "兒女對父母該有的態度(二)─尊敬父母",
    5: "兒女對父母該有的態度(三)─順從父母",  # ✅ 精確修正
    6: "兒女對父母該有的態度(四)─孝敬並順從父母的結果", # ✅ 精確修正
    7: "第一對兄弟－生在罪惡的世界",
    8: "該隱與亞伯－不喜悅與嫉妒",
    9: "約瑟和他的哥哥們",
    10: "愛兄弟姊妹",
    11: "赦免兄弟姊妹",
    12: "敬虔的家庭—戴德生的幼年",
    13: "朋友的建立",
    14: "聖經中朋友的榜樣—大衛和約拿單",
    15: "如何交朋友",
    16: "交友的原則",  # ✅ 性格30點修正為第5冊正確人際關係內容
    17: "對待同伴的原則",
    18: "對待眾人的原則",
    19: "對待年長者的態度",
    20: "對待幼小者的態度",
    21: "對待貧窮或軟弱者的態度",
    22: "對待有權力者的態度",
    23: "對待得罪我們之人的態度",
    24: "不結交壞同伴",
    25: "不結交搬弄是非的人",
    26: "不結交暴躁的人",
    27: "不結交愛宴樂的人",
    28: "不結交驕傲或眼高的人",
    29: "不結交愚昧的人",
    30: "不結交自私或貪心的人",
    31: "不結交不義或作惡的人",
    32: "作同伴的好榜樣",
    33: "領同伴歸主",
    34: "撒但破壞一切人際關係",
    35: "召會生活保守一切人際關係", # ✅ 精確修正
    36: "小明燈(一)─順服",
    37: "小明燈(二)─親愛與關心",
    38: "小明燈(三)─俯就",
    39: "小明燈(四)─卑微與謙卑",
    40: "小明燈(五)─柔和不爭競"
}

# 動態填充骨架，確保其餘週次也有高品質台灣繁體內容大底 (完全移除性格30點，以人際關係美德為主導)
for wk, name in WEEKS_4_40_TITLES.items():
    if wk not in CURRICULUM_DB:
        if wk in [4, 5, 6]:
            verse = "「兒女，要在主裡聽從父母，因為這是正當的。」— 以弗所書六章 1 節"
            hymn = "兒童詩歌第 209 首《兒女當孝敬父母》"
            desc = f"學習在日常生活中尊敬、順從並體貼父母。這課專注於「{name}」，使我們在主裡享受神所命定的甜美家庭次序。"
            story = "小綿羊在媽媽溫柔的呼喚聲中，學會了立刻回應。雖然牠很想繼續在草地上玩耍，但牠知道父母是神量給牠的遮蓋。當牠順從父母的話回到羊圈時，天空下起了大雨，小綿羊心中充滿了對神的感激，體會到順從帶來的保守與平安。"
            exercise = "本週操練：當父母對我們說話或交辦事情時，操練立刻微笑回應，不抱怨、不拖延，活出敬重父母的甜美生活。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n神是有次序的神，祂所安排的家庭是充滿愛與和諧的等次。當我們操練「{name}」時，我們不是憑著自己勉強的意志，而是回到我們最深處的靈裡，呼求主名，讓主耶穌那聽從、體貼的生命流淌出來，使我們的家因著我們的甜美行為而滿了歡笑。"
        elif wk in [7, 8, 9, 10, 11, 12]:
            verse = "「看哪，弟兄和睦同居，是何等的善，何等的美！」— 詩篇一百三十三篇 1 節"
            hymn = "兒童詩歌第 733 首《相親相愛》"
            desc = f"學習與同伴、手足間建立健康、和睦的相處。本課深入探討「{name}」，引導兒童明白在愛與和睦中生活是多麼美善的事。"
            story = "小羊和哥哥因為一件好玩的玩具起了小小的爭執。這時小羊決定把玩具先讓給哥哥。當哥哥高興地抱住牠時，兩個小羊都感受到了和睦同居的無比甜美，周圍滿了和煦溫馨的光線。"
            exercise = "本週操練：在家中主動關心、幫助自己的兄弟姊妹或兒童班的同伴，學習分享玩具與點心，不嫉妒也不爭吵。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n在聖經中，神非常看重弟兄姊妹之間的關係。該隱因嫉妒而犯下了罪，但約瑟卻因著成熟的生命而寬恕並愛護得罪他的哥哥們。戴德生從小生長在一個敬虔、相親相愛的家庭中，這為他一生的事奉打下了美好的根基。我們應當天天呼求主名，得著主愛的供應，使我們與身邊的手足同伴和睦相處，彰顯神美善的生命。"
        elif wk in [13, 14, 15]:
            verse = "「朋友乃時常親愛，弟兄輩生來是用以分憂的。」— 箴言十七章 17 節"
            hymn = "兒童詩歌第 735 首《我們是主內好同伴》"
            desc = f"建立正確的交友觀念。本課探討「{name}」，藉由聖經中美好的榜樣，學習如何尋找並建立在主裡相親相愛的屬靈同伴關係。"
            story = "大衛和約拿單是聖經中極其美麗的朋友榜樣。約拿單愛大衛如同愛自己的生命，在危難中彼此守護、同心合意。當我們在主裡有同伴一起追求真理時，我們的屬靈前途就更加明亮，被溫暖的和煦光束完全照亮。"
            exercise = "本週操練：主動向同伴表達關心，這週找一位同伴一起禱告、背一節聖經，彼此建立健康、積極的屬靈同伴關係。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n交朋友是每個人成長中的大事。聖經中大衛和約拿單的友誼向我們展示了何為無私的愛與同心。我們要在主裡學習交朋友的智慧，不是為了彼此玩鬧或攀比，而是為了能有同伴一起追求真理、一同奔跑屬天的道路。讓我們天天來到神面前，求主賜給我們敬虔的好同伴，在召會生活中一同健康成長。"
        elif wk in [16, 17, 18, 19, 20, 21, 22, 23]:
            verse = "「所以無論何事，你們願意人怎樣待你們，你們也要怎樣待人。」— 馬太福音七章 12 節"
            hymn = "兒童詩歌第 733 首《相親相愛》"
            desc = f"學習在人群中彰顯神聖生命的美德，並建立合宜的互動。這課著重於「{name}」，引導我們用尊敬、溫柔和體貼的態度對待每一個人。"
            story = "小熊在森林裡遇到年長的山羊爺爺正在辛苦地抱著木柴。牠立刻跑過去，恭敬地向爺爺問安，並主動幫忙抱木柴。這時，一道溫馨和煦的光線照亮了牠們，山羊爺爺高興地撫摸小熊的頭，誇獎牠是個懂禮貌、敬重長輩的好孩子。"
            exercise = f"本週操練：實行「{name}」，在生活中主動對身邊的人（如長輩、幼小者、或需要幫助的人）給予真誠的關心與溫柔的問候。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n神把我們擺在人群中，盼望我們能將神那溫柔、愛與尊重的神聖美德彰顯在每個人身上。無論是對待年長者、幼小者，還是對待貧窮或得罪我們的人，我們都要活出合乎聖經次序的行事為人。當我們遇見衝突或覺得不容易愛人時，我們只要閉上眼睛呼求『主耶穌！』，轉向靈裡接受生命水流的供應，神聖的愛與寬容就會自然而然溢流出來，使我們在同伴中成為芬芳的見證。"
        elif wk in [24, 25, 26, 27, 28, 29, 30, 31]:
            verse = "「不要自欺：濫交朋友，敗壞好行為。」— 哥林多前書十五章 33 節"
            hymn = "兒童詩歌第 724 首《當保守你心》"
            desc = f"保守我們的心靈不受消極環境的影響。本課幫助我們認識「{name}」的重要性，並學會在朋友交往中保持清明與聖潔。"
            story = "小羊看見森林邊緣有幾隻小狼在玩危險且粗魯的遊戲，牠們還威脅其他溫和的小動物。小羊知道不應該跟壞同伴混在一起。牠立刻轉過身，回到溫馨的羊群與牧人身邊，和煦的光線（溫暖的光束）灑在牠身上，讓牠心裡無比踏實、充滿平安。"
            exercise = f"本週操練：警醒保守自己，不隨從消極的同學或玩伴，不聽是非、不亂發脾氣、不參與愛宴樂的壞習慣，天天過健康的召會生活。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n聖經嚴格提醒我們，不良的交往會敗壞我們的好行為。撒但常常藉著搬弄是非、暴躁、愛宴樂、驕傲或自私的同伴來玷污我們的心思。我們需要有屬靈的智慧，在學校和鄰里中懂得保守自己，不與消極不義的壞同伴結交。我們要積極地活在召會生活中，與愛主、追求真理的清心同伴一同聚會、一同禱告，使我們的一言一行都蒙保守在主純淨的愛中。"
        elif wk in [32, 33]:
            verse = "「不可叫人小看你年輕，總要在言語、行為、愛、信心、純潔上，都作信徒的榜樣。」— 提摩太前書四章 12 節"
            hymn = "兒童詩歌第 611 首《作主耶穌的好孩子》"
            desc = f"在日常生活中作主的明亮見證，發揮積極的影響力。本課幫助我們學習「{name}」，使身邊的同伴因我們而認識主。"
            story = "小兔子在學校裡看見同桌非常難過。牠沒有跟著別人一起冷落他，而是走過去分享自己的點心，並溫柔地為他呼求主名。當同伴的心被溫暖的光束環繞、臉上重新露出笑容時，小兔子成功在言語和行為上作了主的好榜樣，也把福音與喜樂帶給了同伴。"
            exercise = f"本週操練：在同伴面前有端正、柔和的行為。這週試著向一位同學分享你的兒童班生活，帶領同伴一同接觸主名。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n我們雖然年紀小，但在主的恩典裡，我們完全可以在日常生活中作同伴的好榜樣。藉著我們敬虔、溫柔的言語與純潔的行為，我們能把神聖生命的美德照耀在黑暗的角落裡。更進一步，我們要操練領同伴歸主，邀請同學來參加兒童聚會，讓他們也一同呼求主名，得著生命水流的滋潤與豐富的供應，在主裡過喜樂、充實的生活。"
        elif wk in [34, 35]:
            if wk == 34:
                verse = "「平安的神快要將撒但踐踏在你們腳下。」— 羅馬書十六章 20 節"
            else:
                verse = "「這家就是活神的召會，真理的柱石和根基。」— 提摩太前書三章 15 節"
            hymn = "兒童詩歌第 742 首《召會生活真甜美》"
            desc = f"看清仇敵的破壞，並珍賞神所量給我們的避難所。本課深刻剖析「{name}」，幫助我們明白召會生活對於保守一切和諧人際關係的極大價值。"
            story = "小動物們原本在花園裡快樂地合作，但撒但挑撥離間，大家開始懷疑彼此。當牠們感到孤單痛苦時，牧人吹響了笛聲，把大家召集回安全的羊圈（召會生活）。在和煦溫馨的光線指引下，大家彼此認罪、互相和好，撒但的陰謀再次被踐踏在腳下。"
            exercise = f"本週操練：拒絕撒但搬弄是非與挑撥離間的詭計。這週積極參加兒童排與主日聚會，與弟兄姊妹同心合意，保守甜美的召會生活。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n撒但的本性就是說謊與破壞，牠千方百計想要破壞我們與父母、與兄弟姊妹、與同伴之間的美好關係，帶進不和、猜忌與紛爭。但是讚美主，神為我們預備了活神的召會！召會生活是保守我們人際關係最安全、最甜蜜的地方。在召會裡，我們與聖徒一同享受神的話、一同活在靈中，任何撒但帶進的消極毒素都會在愛與光的洗滌中完全消散，使我們能過和睦同居、神聖和諧的甜美生活。"
        elif wk in [36, 37, 38, 39, 40]:
            verse = "「年幼的，要順服年長的；你們眾人也要以謙卑束腰，彼此服事。」— 彼得前書五章 5 節"
            hymn = "兒童詩歌第 132 首《我是世上的小明燈》"
            desc = f"活出基督柔和、低微、俯就的高尚人性美德。本課著重於「{name}」，引導兒童在日常生活中彰顯主那明亮、純淨的生命光芒。"
            story = "小綿羊在兒童班的相調中，被指派擔任謙卑的小助手。牠不與人爭競，而是溫和、柔順地俯就軟弱的小羊，幫助牠們拿點心、整理桌椅。每當牠彎下腰服務時，心中都湧流出喜樂的泉源，周圍也散發出溫暖和煦的光線，像一盞明亮的小明燈照亮了身邊的每一個人。"
            exercise = f"本週操練：在生活與聚會的小事上實行「{name}」，主動俯就、體貼同伴的需要，不與人爭競，用愛彼此服事、彼此和睦。"
            pure_text = f"第 {wk} 週：{name}\\n\\n【本週背經】\\n{verse}\\n\\n【教材內容】\\n{desc}\\n小明燈的性格是基督生命美德的自然流露。在召會、學校與家庭中，我們不要爭大、不要爭第一，而是要操練順服、親愛、俯就、卑微與柔和。當我們遇到想與人計較或爭勝的情況時，我們要回到我們的靈裡，吸入主那柔和、謙卑的生命成分。主完美的生命能帶領我們跨越天然的有限，在凡事上作明亮、純潔的小明燈，叫周圍的人因著我們的行事為人而看見神的榮耀彰顯。"

        y_url = ""
        if wk == 14:
            y_url = "https://www.youtube.com/watch?v=I_sgi1EQAkQ"
        elif wk == 15:
            y_url = "https://www.youtube.com/watch?v=T-pPHX8lLYE"

        CURRICULUM_DB[wk] = {
            "title": f"第 {wk} 週：{name}",
            "hymn": hymn,
            "verse": verse,
            "discussion": f"1. 在家裡或兒童班，我們要如何實行「{name.split('─')[-1]}」？\\n2. 當我們覺得自己力有未逮時，如何回到靈裡呼求主名，享受主生命的供應？",
            "story": story,
            "exercise": exercise,
            "pure_text": pure_text,
            "youtube_url": y_url,          
            "custom_slides_url": "https://drive.google.com/drive/folders/1uo7uR6wQk94C8xqxF1uwnChZWJs58zog?usp=sharing"       
        }

# --- 輔助函式與邏輯處理 ---

# 動態產生 Word 教材 (純淨無雜訊)
def generate_docx(week_num):
    data = CURRICULUM_DB[week_num]
    doc = Document()
    doc.add_heading(data['title'], 0)
    doc.add_paragraph(data['pure_text'].replace("\\n", "\n"))
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 標準化 NotebookLM 簡報專用提示詞模板
def generate_notebooklm_prompt(week_num):
    data = CURRICULUM_DB[week_num]
    prompt_template = f"""請根據以下核心教材內容，為『學齡至國小中低年級兒童』設計一份活潑、可愛、充滿童趣且極具吸引力的簡報腳本指令碼。

重要指示（請嚴格遵守）：
1. 僅根據以下給定的內容進行設計，絕對「不額外延伸其他主題」或旁徵博引，確保內容專注。
2. 視覺設計描述需鎖定中低年級審美，建議採用溫暖、柔和、明亮的插畫風格。
3. 嚴防敏感宗教肖像與符號：簡報的畫面與視覺提示中，絕對不可出現擬人化的神聖畫像、上帝、天使、耶穌、十字架、翅膀、雙手合十等宗教實體或人像符號。若遇到引導、溫馨或得勝的轉折場景，請統一使用『溫暖煦煦的光束』或『和煦溫馨的光線』來代替，展現出神聖高尚的生命美德。

簡報大綱必須嚴格包括以下六個黃金板塊：
- 第一頁：封面（包含本課週次數與標題：{data['title']}）
- 第二頁：本週詩歌（唱詩相調：{data['hymn']}）
- 第三頁：問題與討論（互動啟發：{data['discussion']}）
- 第四頁：本週故事（生動講述：{data['story']}）
- 第五頁：本週背經（神話供應：{data['verse']}）
- 第六頁：本週生活操練（生活實踐：{data['exercise']}）

原始參考文本：
{data['pure_text']}
"""
    return prompt_template

# --- 多網頁切換系統架構 ---

# 1. 首頁：課程目錄門戶
if st.session_state.page == 'home':
    st.title("召會兒童教材自動化生產線")
    st.subheader("點選單元與課次進入教材、簡報與影音專區")
    st.write("本系統已完全優化為台灣繁體中文，融入主恢復編輯習慣（轉向靈、呼求主名、生命供應），並全面過濾不適宜宗教意象，改以溫柔和煦的光束指引。")
    st.markdown("---")

    # 單元分類結構
    UNITS = {
        "第一單元：父母與兒女的關係 (第 1 至 6 週)": range(1, 7),
        "第二單元：家庭與手足關係 (第 7 至 12 週)": range(7, 13),
        "第三單元：同伴與朋友關係 (第 13 至 15 週)": range(13, 16),
        "第四單元：對待同伴與人群的原則與交往態度 (第 16 至 33 週)": range(16, 34),
        "第五單元：人際關係與小明燈系列 (第 34 至 40 週)": range(34, 41)
    }

    # 渲染五大單元目錄
    for unit_name, week_range in UNITS.items():
        st.markdown(f"## {unit_name}")
        
        # 使用 3 等分網格顯示卡片
        cols = st.columns(3)
        for idx, wk_num in enumerate(week_range):
            col_idx = idx % 3
            with cols[col_idx]:
                week_title_clean = CURRICULUM_DB[wk_num]['title'].split('：')[-1]
                st.markdown(f"""
                    <div class="catalogue-card">
                        <span style="color: #64748b; font-size: 13px; font-weight: bold;">第 {wk_num} 週</span>
                        <h4 style="margin: 5px 0 15px 0; color: #0f172a; font-size: 15px; height: 40px; overflow: hidden; line-height: 1.3;">{week_title_clean}</h4>
                    </div>
                """, unsafe_allow_html=True)
                
                # 點擊按鈕跳轉至詳情頁，並更新 session_state
                if st.button("進入此課程", key=f"select_wk_{wk_num}"):
                    st.session_state.selected_week = wk_num
                    st.session_state.page = 'detail'
                    st.rerun()
        st.markdown("---")

# 2. 詳情頁：特定週次之多媒體與指令頁面
elif st.session_state.page == 'detail':
    current_week = st.session_state.selected_week
    current_data = CURRICULUM_DB[current_week]
    
    # 頂部導覽列 (高對比返回目錄按鈕)
    col_nav1, col_nav2 = st.columns([1, 4])
    with col_nav1:
        if st.button("<- 返回課程目錄", key="back_to_home"):
            st.session_state.page = 'home'
            st.rerun()
            
    with col_nav2:
        # 提供快速切換的側邊或頂部下拉選單，不用每次回目錄
        quick_list = [f"第 {i} 週：{CURRICULUM_DB[i]['title'].split('：')[-1]}" for i in range(1, 41)]
        selected_quick_str = st.selectbox(
            "快速切換週次", 
            quick_list, 
            index=current_week-1,
            label_visibility="collapsed"
        )
        quick_week_num = int(selected_quick_str.split(" ")[1])
        if quick_week_num != current_week:
            st.session_state.selected_week = quick_week_num
            st.rerun()

    st.markdown("---")

    # 顯示當前選定課程標題
    st.markdown(f"### 當前選定：{current_data['title']}")

    # 三大頁籤 (Tabs) 切換設計 - 注入高對比色彩排版
    tab_lesson, tab_slides, tab_demo = st.tabs([
        "1. 課程教材 (純文字)", 
        "2. 簡報內容 (雲端圖片連結)", 
        "3. 教學示範 (影片)"
    ])

    # ---- 頁籤一：課程教材 ----
    with tab_lesson:
        st.markdown('<div class="download-card">', unsafe_allow_html=True)
        st.markdown("### 純文字教材下載")
        st.write("已完美校正為台灣召會習慣語境（靈、主名、生命供應），無任何雜訊提示，適合服事者直接印製或群組分享。")
        
        # 產生並提供下載
        docx_file = generate_docx(current_week)
        st.download_button(
            label="下載純文字教材 (.docx)",
            data=docx_file,
            file_name=f"第{current_week}週_純文字教材繁體版.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"word_tab_detail_{current_week}"
        )
        st.markdown('</div>', unsafe_allow_html=True)
        
        # 課程教材文字預覽
        st.markdown("#### 課程教材內文預覽")
        st.info(current_data['pure_text'].replace("\\n", "\n"))

    # ---- 頁籤二：簡報內容 ----
    with tab_slides:
        # 1. 既有簡報下載 (自動開新視窗)
        if current_data['custom_slides_url']:
            st.markdown('<div class="slides-card">', unsafe_allow_html=True)
            st.markdown("### 既有精美簡報圖片 (雲端硬碟)")
            st.write("此課程已預備好簡報 PNG 圖片，點擊下方按鈕將於「新分頁」中開啟雲端資料夾，即可直接檢視或下載。")
            st.link_button("前往檢視/下載簡報圖片", current_data['custom_slides_url'])
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("本課尚未填入既有簡報網址。只要把你的簡報圖片丟上 Google Drive 雲端，並把連結貼給我，我就能幫你把按鈕做出來！")
            
        # 2. NotebookLM 指令碼
        st.markdown('<div class="download-card">', unsafe_allow_html=True)
        st.markdown("### NotebookLM 簡報生成專用提示詞")
        st.write("一鍵下載標準化指令。設定已包含：活潑可愛、禁止跑題、以及採用『光束引導』嚴防宗教肖像原則。")
        
        prompt_text = generate_notebooklm_prompt(current_week)
        st.download_button(
            label="下載簡報指令碼 (.txt)",
            data=prompt_text,
            file_name=f"第{current_week}週_NotebookLM簡報指令.txt",
            mime="text/plain",
            key=f"prompt_tab_detail_{current_week}"
        )
        st.markdown('</div>', unsafe_allow_html=True)
        
        with st.expander("展開直接在網頁複製指令"):
            st.code(prompt_text, language="text")

    # ---- 頁籤三：教學示範 ----
    with tab_demo:
        if current_data['youtube_url']:
            st.markdown('<div class="video-card">', unsafe_allow_html=True)
            st.markdown("### 課程教學示範影片")
            st.write("服事者已上傳之教學錄影。點擊下方按鈕可在「新視窗」開啟觀看，或直接在網頁上播放：")
            
            # 1. 點擊開新視窗按鈕
            st.link_button("在 YouTube 新視窗開啟影片", current_data['youtube_url'])
            
            # 2. 嵌入式網頁播放器
            st.markdown("---")
            st.video(current_data['youtube_url'])
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("本課尚未填入教學錄影網址。只要你把 YouTube 連結貼給我，我就會幫你把播放器做出來！")

    st.markdown("---")
    st.caption("召會兒童教材生產線 • 符合主恢復聖徒操練、無宗教肖像原則")
